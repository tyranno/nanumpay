#!/usr/bin/env python3
"""
DOCX 파일의 모든 테이블에 테두리를 추가하는 스크립트

사용법:
1. python3-docx 설치:
   sudo apt install python3-docx -y

2. 스크립트 실행:
   python3 add_table_borders.py
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_borders(table):
    """테이블에 모든 테두리 추가"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 기존 테두리 제거
    for child in list(tblPr):
        if child.tag == qn('w:tblBorders'):
            tblPr.remove(child)

    # 새 테두리 설정
    tblBorders = OxmlElement('w:tblBorders')

    # 모든 테두리 타입에 대해 설정
    border_names = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    for border_name in border_names:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 단선
        border.set(qn('w:sz'), '4')        # 두께 (1/8 pt 단위, 4 = 0.5pt)
        border.set(qn('w:space'), '0')     # 간격
        border.set(qn('w:color'), '000000') # 검은색
        tblBorders.append(border)

    tblPr.append(tblBorders)

def main():
    docx_path = 'docs/시스템_요구사항_검토문서.docx'

    print(f'DOCX 파일 로드 중: {docx_path}')
    doc = Document(docx_path)

    print(f'총 {len(doc.tables)}개의 테이블 발견')

    # 모든 테이블에 테두리 추가
    for idx, table in enumerate(doc.tables, 1):
        add_table_borders(table)
        print(f'  테이블 {idx}/{len(doc.tables)} 처리 완료')

    print(f'파일 저장 중: {docx_path}')
    doc.save(docx_path)

    print('✅ 완료! 모든 테이블에 테두리가 추가되었습니다.')

if __name__ == '__main__':
    main()
